#!/usr/bin/env python3
"""
md2docx.py — Convert a structured Markdown BRD into a professionally formatted .docx.

Supports: ATX headings (#..######), pipe tables (with header row + --- separator),
ordered/unordered lists (nested by indent), bold (**x**) and inline code (`x`),
fenced code blocks (```), blockquotes (>), horizontal rules (---), and a generated
Table of Contents field. Applies the BRD house style: dark-blue headings, light-blue
shaded table header rows, Arial body, page numbers in footer, title + "Confidential"
in header.

Usage: python3 md2docx.py <input.md> <output.docx> ["Document Title"]
"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MID_BLUE = RGBColor(0x2E, 0x5B, 0x9E)
HEADER_SHADE = "DCE6F1"   # light blue for table header row
BODY_FONT = "Arial"

HEADING_SIZES = {1: 18, 2: 15, 3: 13, 4: 12, 5: 11, 6: 11}


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_code
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = ""
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instrText, fldChar2, t, fldChar3):
        run._r.append(el)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline_runs(paragraph, text):
    """Render bold (**x**) and inline code (`x`) inside a paragraph."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def style_body(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def setup_header_footer(doc, title):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = title + "\tConfidential"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Page ").font.size = Pt(9)
    add_field(fp, "PAGE")
    fp.add_run(" of ").font.size = Pt(9)
    add_field(fp, "NUMPAGES")


def parse_table(lines, i):
    """Parse a pipe table starting at line i. Returns (rows, next_i)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    cells = []
    for r in rows:
        parts = [c.strip() for c in r.strip().strip("|").split("|")]
        cells.append(parts)
    # drop separator row (---|---)
    if len(cells) >= 2 and all(set(c) <= set("-: ") for c in cells[1]):
        header = cells[0]
        body = cells[2:]
    else:
        header = cells[0]
        body = cells[1:]
    return header, body, i


def add_table(doc, header, body):
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j, htext in enumerate(header):
        set_cell_background(hdr[j], HEADER_SHADE)
        p = hdr[j].paragraphs[0]
        add_inline_runs(p, htext)
        for run in p.runs:
            run.bold = True; run.font.size = Pt(10); run.font.color.rgb = DARK_BLUE
    for row in body:
        cells = table.add_row().cells
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            p = cells[j].paragraphs[0]
            add_inline_runs(p, val)
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def convert(md_path, docx_path, title=None):
    with open(md_path, "r", encoding="utf-8") as f:
        raw = f.read()
    lines = raw.split("\n")

    doc = Document()
    style_body(doc)

    # Derive title from first H1 if not given
    if title is None:
        m = re.search(r"^#\s+(.+)$", raw, re.M)
        title = m.group(1).strip() if m else "Business Requirements Document"

    setup_header_footer(doc, title)

    # Title page
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = DARK_BLUE
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Business Requirements Document"); r.font.size = Pt(14); r.font.color.rgb = MID_BLUE
    cf = doc.add_paragraph(); cf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cf.add_run("CONFIDENTIAL"); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_page_break()

    # TOC
    th = doc.add_paragraph(); r = th.add_run("Table of Contents"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DARK_BLUE
    tocp = doc.add_paragraph()
    add_field(tocp, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    i = 0
    n = len(lines)
    first_h1_skipped = False
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            set_cell_background_para(p, "F2F2F2")
            continue

        # Tables
        if stripped.startswith("|"):
            header, body, i = parse_table(lines, i)
            add_table(doc, header, body)
            continue

        # Headings
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                i += 1
                continue  # already used as title page
            p = doc.add_heading(level=min(level, 4))
            p.runs.clear() if p.runs else None
            run = p.add_run(text)
            run.font.color.rgb = DARK_BLUE if level <= 2 else MID_BLUE
            run.font.size = Pt(HEADING_SIZES.get(level, 11))
            run.bold = True
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^(\*\s*){3,}$", stripped) or re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline_runs(p, stripped.lstrip(">").strip())
            i += 1
            continue

        # Lists
        lm = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if lm:
            indent = len(lm.group(1))
            marker = lm.group(2)
            text = lm.group(3)
            ordered = bool(re.match(r"\d+\.", marker))
            style = "List Number" if ordered else "List Bullet"
            p = doc.add_paragraph(style=style)
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            add_inline_runs(p, text)
            i += 1
            continue

        # Blank
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def set_cell_background_para(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 md2docx.py <input.md> <output.docx> [\"Title\"]")
        sys.exit(1)
    title = sys.argv[3] if len(sys.argv) > 3 else None
    convert(sys.argv[1], sys.argv[2], title)
